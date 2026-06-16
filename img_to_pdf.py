# img_to_pdf.py
# VERSAO_VISIVEL: v2.2 - inclui verificação Office e modo simples sem Office
import csv
import io
import sys
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from xml.sax.saxutils import escape

import img2pdf
from PIL import Image, ImageDraw, ImageFont, ImageTk
from pypdf import PdfReader, PdfWriter
from tkinterdnd2 import TkinterDnD, DND_FILES


IMG_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".gif"}
WORD_EXTS = {".doc", ".docx", ".docm", ".dot", ".dotx", ".dotm", ".rtf"}
EXCEL_EXTS = {".xls", ".xlsx", ".xlsm", ".xlsb", ".xlt", ".xltx", ".xltm", ".csv"}
OFFICE_EXTS = WORD_EXTS | EXCEL_EXTS

# Formatos que o modo simples consegue ler sem Word/Excel/LibreOffice.
SIMPLE_WORD_EXTS = {".docx", ".docm"}
SIMPLE_EXCEL_EXTS = {".xlsx", ".xlsm", ".xltx", ".xltm", ".csv"}

A4_W_PT = 595.276
A4_H_PT = 841.89
MARGIN_PT = 18


def _safe_pdf_path(path_str: str) -> Path:
    p = Path(path_str.strip().strip('"'))
    if p.suffix.lower() != ".pdf":
        p = p.with_suffix(".pdf")
    return p


def _parse_dnd_files(data: str) -> list[str]:
    # Windows costuma enviar: "{C:\\path com espaço\\1.pdf} {C:\\outro\\2.pdf}"
    out, cur, in_brace = [], "", False
    for ch in data:
        if ch == "{":
            in_brace = True
            cur = ""
            continue
        if ch == "}" and in_brace:
            in_brace = False
            if cur:
                out.append(cur)
            cur = ""
            continue
        if ch == " " and not in_brace:
            if cur:
                out.append(cur)
                cur = ""
            continue
        cur += ch
    if cur:
        out.append(cur)
    return [p.strip().strip('"') for p in out if p.strip()]


def _dpi_tuple(dpi) -> tuple[float, float]:
    if dpi is None:
        return 96.0, 96.0
    if isinstance(dpi, (tuple, list)) and len(dpi) >= 2:
        x = float(dpi[0] or 96.0)
        y = float(dpi[1] or 96.0)
        return (x if x > 0 else 96.0), (y if y > 0 else 96.0)
    try:
        d = float(dpi)
        return (d if d > 0 else 96.0), (d if d > 0 else 96.0)
    except Exception:
        return 96.0, 96.0


def a4_fit_layout_fun(imgw_px, imgh_px, dpi):
    dx, dy = _dpi_tuple(dpi)
    imgw_pt = (float(imgw_px) / dx) * 72.0
    imgh_pt = (float(imgh_px) / dy) * 72.0

    if imgw_pt > imgh_pt:
        page_w, page_h = A4_H_PT, A4_W_PT
    else:
        page_w, page_h = A4_W_PT, A4_H_PT

    max_w = max(1.0, page_w - 2 * MARGIN_PT)
    max_h = max(1.0, page_h - 2 * MARGIN_PT)
    scale = min(max_w / imgw_pt, max_h / imgh_pt)
    return page_w, page_h, imgw_pt * scale, imgh_pt * scale


# ===================== Imagens -> PDF =====================


def convert_image_to_pdf_single(input_image: Path, output_pdf: Path, fit_a4: bool) -> None:
    if not input_image.exists() or not input_image.is_file():
        raise FileNotFoundError(f"Imagem não encontrada: {input_image}")

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    layout = a4_fit_layout_fun if fit_a4 else None

    # 1) tenta embutir sem recompressão.
    try:
        pdf_bytes = img2pdf.convert(str(input_image), layout_fun=layout) if layout else img2pdf.convert(str(input_image))
        output_pdf.write_bytes(pdf_bytes)
        return
    except Exception:
        pass

    # 2) fallback: PNG lossless em memória.
    try:
        with Image.open(input_image) as im:
            buf = io.BytesIO()
            im.save(buf, format="PNG", optimize=False)
            png_bytes = buf.getvalue()
        pdf_bytes = img2pdf.convert(png_bytes, layout_fun=layout) if layout else img2pdf.convert(png_bytes)
        output_pdf.write_bytes(pdf_bytes)
    except Exception as e:
        raise RuntimeError(f"Falha ao converter a imagem para PDF: {e}") from e


def convert_images_to_one_pdf(image_paths: list[Path], output_pdf: Path, fit_a4: bool) -> None:
    if not image_paths:
        raise ValueError("Nenhuma imagem foi selecionada.")

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    layout = a4_fit_layout_fun if fit_a4 else None

    # 1) tenta direto (melhor para evitar recompressão).
    try:
        pdf_bytes = img2pdf.convert([str(p) for p in image_paths], layout_fun=layout) if layout else img2pdf.convert([str(p) for p in image_paths])
        output_pdf.write_bytes(pdf_bytes)
        return
    except Exception:
        pass

    # 2) fallback: converte tudo para PNG lossless em memória.
    try:
        payloads: list[bytes] = []
        for p in image_paths:
            with Image.open(p) as im:
                buf = io.BytesIO()
                im.save(buf, format="PNG", optimize=False)
                payloads.append(buf.getvalue())
        pdf_bytes = img2pdf.convert(payloads, layout_fun=layout) if layout else img2pdf.convert(payloads)
        output_pdf.write_bytes(pdf_bytes)
    except Exception as e:
        raise RuntimeError(f"Falha ao converter imagens para PDF: {e}") from e


# ===================== Juntar PDFs =====================


def merge_pdfs(input_pdfs: list[Path], output_pdf: Path, password_cb) -> None:
    if not input_pdfs:
        raise ValueError("Nenhum PDF foi selecionado.")

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    writer = PdfWriter()

    for pdf_path in input_pdfs:
        if not pdf_path.exists() or not pdf_path.is_file():
            raise FileNotFoundError(f"PDF não encontrado: {pdf_path}")

        reader = PdfReader(str(pdf_path))

        if reader.is_encrypted:
            ok = False
            try:
                ok = bool(reader.decrypt(""))
            except Exception:
                ok = False

            if not ok:
                pwd = password_cb(pdf_path.name)
                if not pwd:
                    raise RuntimeError(f"Senha não informada para: {pdf_path.name}")
                try:
                    if not reader.decrypt(pwd):
                        raise RuntimeError(f"Senha incorreta para: {pdf_path.name}")
                except Exception as e:
                    raise RuntimeError(f"Não foi possível abrir (criptografado): {pdf_path.name}. Detalhe: {e}") from e

        for page in reader.pages:
            writer.add_page(page)

    with open(output_pdf, "wb") as f:
        writer.write(f)


# ===================== Word/Excel -> PDF =====================


def detect_ms_office() -> dict:
    """
    Retorna se Word/Excel estão disponíveis via automação COM.
    Só funciona em Windows. Em outros sistemas, retorna indisponível.
    """
    result = {
        "platform_ok": sys.platform.startswith("win"),
        "pywin32_ok": False,
        "word_ok": False,
        "excel_ok": False,
        "word_error": "",
        "excel_error": "",
        "general_error": "",
    }

    if not result["platform_ok"]:
        result["general_error"] = "Automação do Microsoft Office só está disponível no Windows."
        return result

    try:
        import pythoncom
        import win32com.client
        result["pywin32_ok"] = True
    except Exception as e:
        result["general_error"] = f"pywin32 não está instalado ou não carregou corretamente: {e}"
        return result

    pythoncom.CoInitialize()
    try:
        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            result["word_ok"] = True
        except Exception as e:
            result["word_error"] = str(e)
        finally:
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass

        excel = None
        try:
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            result["excel_ok"] = True
        except Exception as e:
            result["excel_error"] = str(e)
        finally:
            try:
                if excel is not None:
                    excel.Quit()
            except Exception:
                pass

    finally:
        pythoncom.CoUninitialize()

    return result


def convert_word_with_ms_office(src: Path, out_pdf: Path) -> None:
    try:
        import pythoncom
        import win32com.client
    except ImportError as e:
        raise RuntimeError("Dependência ausente: instale pywin32 com: pip install pywin32") from e

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(src), ReadOnly=True, AddToRecentFiles=False)
        doc.ExportAsFixedFormat(str(out_pdf), 17)  # 17 = wdExportFormatPDF
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()


def convert_excel_with_ms_office(src: Path, out_pdf: Path) -> None:
    try:
        import pythoncom
        import win32com.client
    except ImportError as e:
        raise RuntimeError("Dependência ausente: instale pywin32 com: pip install pywin32") from e

    pythoncom.CoInitialize()
    excel = None
    wb = None
    try:
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        wb = excel.Workbooks.Open(str(src), ReadOnly=True)
        wb.ExportAsFixedFormat(0, str(out_pdf))  # 0 = xlTypePDF
    finally:
        try:
            if wb is not None:
                wb.Close(False)
        except Exception:
            pass
        try:
            if excel is not None:
                excel.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()


def _fmt_cell(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value)
    return str(value)


def _truncate_text(text: str, max_len: int = 300) -> str:
    text = text.replace("\x00", "").strip()
    return text if len(text) <= max_len else text[: max_len - 3] + "..."


def convert_word_simple_to_pdf(src: Path, out_pdf: Path) -> None:
    """
    Conversão simples sem Word/LibreOffice.
    Lê DOCX/DOCM e recria um PDF básico. Não preserva layout fiel.
    """
    if src.suffix.lower() not in SIMPLE_WORD_EXTS:
        raise RuntimeError(
            f"Modo simples não suporta '{src.suffix}'. Para .doc/.rtf/.dot, use o modo Office fiel ou LibreOffice."
        )

    try:
        from docx import Document
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    except ImportError as e:
        raise RuntimeError(
            "Dependências ausentes para modo simples. Instale: pip install python-docx reportlab"
        ) from e

    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    normal = styles["BodyText"]
    heading = styles["Heading2"]

    story = [Paragraph(escape(src.name), heading), Spacer(1, 8)]
    document = Document(str(src))

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            story.append(Paragraph(escape(_truncate_text(text, 2000)), normal))
            story.append(Spacer(1, 5))

    for table in document.tables:
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                row_data.append(Paragraph(escape(_truncate_text(cell_text, 500)), normal))
            if any(str(c.getPlainText()).strip() for c in row_data):
                data.append(row_data)

        if data:
            col_count = max(len(r) for r in data)
            for r in data:
                while len(r) < col_count:
                    r.append(Paragraph("", normal))
            tbl = Table(data, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]))
            story.append(Spacer(1, 8))
            story.append(tbl)
            story.append(Spacer(1, 8))

    if len(story) <= 2:
        story.append(Paragraph("Documento sem texto/tabela extraível pelo modo simples.", normal))

    doc_pdf = SimpleDocTemplate(str(out_pdf), pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    doc_pdf.build(story)


def _excel_rows_from_csv(src: Path) -> list[list[str]]:
    raw = src.read_text(encoding="utf-8-sig", errors="replace")
    sample = raw[:2048]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=";,\t,")
    except Exception:
        dialect = csv.excel
        dialect.delimiter = ";" if ";" in sample and "," not in sample else ","
    rows = []
    for row in csv.reader(io.StringIO(raw), dialect):
        rows.append([_truncate_text(_fmt_cell(v), 300) for v in row])
    return rows


def _add_table_to_story(story, rows: list[list[str]], title: str, styles, landscape_mode=True) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, PageBreak

    if not rows:
        story.append(Paragraph(escape(title), styles["Heading2"]))
        story.append(Paragraph("Sem dados extraíveis.", styles["BodyText"]))
        story.append(Spacer(1, 8))
        return

    # Remove linhas totalmente vazias e normaliza a quantidade de colunas.
    cleaned = []
    for row in rows:
        vals = [str(v or "") for v in row]
        while vals and vals[-1] == "":
            vals.pop()
        if any(v.strip() for v in vals):
            cleaned.append(vals)

    if not cleaned:
        story.append(Paragraph(escape(title), styles["Heading2"]))
        story.append(Paragraph("Sem dados preenchidos.", styles["BodyText"]))
        story.append(Spacer(1, 8))
        return

    max_cols = max(len(r) for r in cleaned)
    for r in cleaned:
        while len(r) < max_cols:
            r.append("")

    # Evita tabelas monstruosas quebrando o PDF. Ainda exporta todas as linhas, mas truncando texto de célula.
    data = []
    for row in cleaned:
        data.append([Paragraph(escape(_truncate_text(v, 250)), styles["BodyText"]) for v in row])

    story.append(Paragraph(escape(title), styles["Heading2"]))
    page_w, _ = landscape(A4) if landscape_mode else A4
    content_w = page_w - 72
    col_w = max(28, content_w / max_cols)
    tbl = Table(data, colWidths=[col_w] * max_cols, repeatRows=1 if len(data) > 1 else 0)
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 2),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(tbl)
    story.append(PageBreak())


def convert_excel_simple_to_pdf(src: Path, out_pdf: Path) -> None:
    """
    Conversão simples sem Excel/LibreOffice.
    Lê XLSX/XLSM/XLTX/XLTM/CSV e recria um PDF básico. Não preserva layout fiel.
    """
    ext = src.suffix.lower()
    if ext not in SIMPLE_EXCEL_EXTS:
        raise RuntimeError(
            f"Modo simples não suporta '{src.suffix}'. Para .xls/.xlsb/.xlt, use o modo Office fiel ou LibreOffice."
        )

    try:
        from openpyxl import load_workbook
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError as e:
        raise RuntimeError(
            "Dependências ausentes para modo simples. Instale: pip install openpyxl reportlab"
        ) from e

    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    story = [Paragraph(escape(src.name), styles["Title"]), Spacer(1, 8)]

    if ext == ".csv":
        rows = _excel_rows_from_csv(src)
        _add_table_to_story(story, rows, "CSV", styles, landscape_mode=True)
    else:
        wb = load_workbook(str(src), data_only=True, read_only=True)
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([_truncate_text(_fmt_cell(v), 300) for v in row])
            _add_table_to_story(story, rows, ws.title, styles, landscape_mode=True)

    if len(story) <= 2:
        story.append(Paragraph("Arquivo sem dados extraíveis pelo modo simples.", styles["BodyText"]))

    doc_pdf = SimpleDocTemplate(
        str(out_pdf),
        pagesize=landscape(A4),
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36,
    )
    doc_pdf.build(story)


def convert_office_file_simple(src: Path, out_pdf: Path) -> None:
    ext = src.suffix.lower()
    if ext in WORD_EXTS:
        convert_word_simple_to_pdf(src, out_pdf)
    elif ext in EXCEL_EXTS:
        convert_excel_simple_to_pdf(src, out_pdf)
    else:
        raise ValueError(f"Extensão não suportada: {src.name}")


def convert_office_file_ms(src: Path, out_pdf: Path) -> None:
    ext = src.suffix.lower()
    if ext in WORD_EXTS:
        convert_word_with_ms_office(src, out_pdf)
    elif ext in EXCEL_EXTS:
        convert_excel_with_ms_office(src, out_pdf)
    else:
        raise ValueError(f"Extensão não suportada: {src.name}")


def convert_office_files_to_pdf(
    files: list[Path],
    out_dir: Path,
    unique_cb,
    status_cb=None,
    mode: str = "auto",
    office_status: dict | None = None,
) -> tuple[list[Path], list[str]]:
    """
    mode:
      - auto: usa Office quando disponível para o tipo do arquivo; senão tenta modo simples.
      - office: exige Word/Excel instalado.
      - simple: não usa Office; recria PDF básico.
    """
    if not files:
        raise ValueError("Nenhum arquivo Word/Excel foi selecionado.")

    out_dir.mkdir(parents=True, exist_ok=True)
    converted: list[Path] = []
    errors: list[str] = []
    office_status = office_status or {}

    for idx, src in enumerate(files, start=1):
        if not src.exists() or not src.is_file():
            errors.append(f"{src.name}: arquivo não encontrado.")
            continue

        ext = src.suffix.lower()
        out_pdf = unique_cb(out_dir / f"{src.stem}.pdf")

        try:
            chosen_mode = mode
            if mode == "auto":
                if ext in WORD_EXTS and office_status.get("word_ok"):
                    chosen_mode = "office"
                elif ext in EXCEL_EXTS and office_status.get("excel_ok"):
                    chosen_mode = "office"
                else:
                    chosen_mode = "simple"

            if status_cb:
                if chosen_mode == "office":
                    status_cb(f"Convertendo {idx}/{len(files)} com Office: {src.name}")
                else:
                    status_cb(f"Convertendo {idx}/{len(files)} no modo simples: {src.name}")

            if chosen_mode == "office":
                convert_office_file_ms(src, out_pdf)
            elif chosen_mode == "simple":
                convert_office_file_simple(src, out_pdf)
            else:
                raise RuntimeError(f"Modo inválido: {mode}")

            converted.append(out_pdf)

        except Exception as e:
            errors.append(f"{src.name}: {e}")

    if status_cb:
        status_cb("Conversão concluída." if not errors else "Conversão concluída com erros.")

    return converted, errors


class App(TkinterDnD.Tk):
    def __init__(self):
        super().__init__()
        self.title("Ferramentas PDF (Tkinter) - v2.2")
        self._set_app_icon()
        self.geometry("960x650")
        self.minsize(960, 650)

        self._merge_seen = set()
        self._img_seen = set()
        self._office_seen = set()
        self._office_status = {}

        self._build_ui()

    def _generate_pdf_icon_png(self, size: int = 64) -> Image.Image:
        im = Image.new("RGBA", (size, size), (196, 0, 0, 255))
        draw = ImageDraw.Draw(im)
        fold = size // 4
        draw.polygon([(size - fold, 0), (size, 0), (size, fold)], fill=(255, 255, 255, 220))
        text = "PDF"
        font = ImageFont.load_default()
        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.text(((size - tw) // 2, (size - th) // 2), text, font=font, fill=(255, 255, 255, 255))
        return im

    def _set_app_icon(self):
        base = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
        ico = base / "pdf.ico"
        png = base / "pdf.png"
        try:
            if ico.exists():
                self.iconbitmap(str(ico))
                return
        except Exception:
            pass
        try:
            im = Image.open(png).convert("RGBA") if png.exists() else self._generate_pdf_icon_png(64)
            self._tk_icon = ImageTk.PhotoImage(im)
            self.iconphoto(True, self._tk_icon)
        except Exception:
            pass

    def _build_ui(self):
        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True, padx=10, pady=10)

        tab_img = ttk.Frame(nb)
        tab_merge = ttk.Frame(nb)
        tab_office = ttk.Frame(nb)

        nb.add(tab_img, text="Imagem → PDF")
        nb.add(tab_merge, text="Juntar PDFs")
        nb.add(tab_office, text="Word/Excel → PDF")

        self._build_tab_image(tab_img)
        self._build_tab_merge(tab_merge)
        self._build_tab_office(tab_office)

    def _normalize_path(self, p: Path) -> str:
        try:
            return str(p.resolve())
        except Exception:
            return str(p)

    def _unique_outfile(self, base_path: Path) -> Path:
        if not base_path.exists():
            return base_path
        stem, suf = base_path.stem, base_path.suffix
        i = 2
        while True:
            candidate = base_path.with_name(f"{stem}_{i}{suf}")
            if not candidate.exists():
                return candidate
            i += 1

    # ===================== Aba 1: Imagens -> PDF(s) =====================

    def _build_tab_image(self, parent: ttk.Frame):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        top = ttk.Frame(parent)
        top.grid(row=0, column=0, sticky="ew", padx=8, pady=(12, 6))
        top.columnconfigure(3, weight=1)

        ttk.Button(top, text="Adicionar imagens...", command=self._add_images).grid(row=0, column=0, sticky="w")
        ttk.Button(top, text="Remover selecionada", command=self._remove_selected_image).grid(row=0, column=1, padx=8)
        ttk.Button(top, text="Limpar lista", command=self._clear_images).grid(row=0, column=2)

        self.one_pdf_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(top, text="PDF único", variable=self.one_pdf_var, command=self._on_one_pdf_toggle).grid(row=0, column=4, sticky="e", padx=(8, 0))

        self.fit_a4_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(top, text="Ajustar tamanho (A4)", variable=self.fit_a4_var).grid(row=0, column=5, sticky="e", padx=(10, 0))

        mid = ttk.LabelFrame(parent, text="Arraste e solte imagens aqui (a ordem da lista será a ordem do PDF)")
        mid.grid(row=1, column=0, sticky="nsew", padx=8, pady=6)
        mid.columnconfigure(0, weight=1)
        mid.rowconfigure(0, weight=1)

        self.img_list: list[Path] = []
        self.img_listbox = tk.Listbox(mid, selectmode=tk.SINGLE)
        self.img_listbox.grid(row=0, column=0, sticky="nsew", padx=(8, 0), pady=8)

        sb = ttk.Scrollbar(mid, orient="vertical", command=self.img_listbox.yview)
        sb.grid(row=0, column=1, sticky="ns", pady=8)
        self.img_listbox.configure(yscrollcommand=sb.set)

        side = ttk.Frame(mid)
        side.grid(row=0, column=2, sticky="ns", padx=10, pady=8)
        ttk.Button(side, text="↑ Subir", command=self._img_move_up).grid(row=0, column=0, pady=(0, 6), sticky="ew")
        ttk.Button(side, text="↓ Descer", command=self._img_move_down).grid(row=1, column=0, pady=(0, 6), sticky="ew")

        for w in (mid, self.img_listbox):
            try:
                w.drop_target_register(DND_FILES)
                w.dnd_bind("<<Drop>>", self._on_drop_images)
            except Exception:
                pass

        bottom = ttk.Frame(parent)
        bottom.grid(row=2, column=0, sticky="ew", padx=8, pady=(10, 6))
        bottom.columnconfigure(1, weight=1)

        self.img_out_label = ttk.Label(bottom, text="Salvar PDF final em:")
        self.img_out_label.grid(row=0, column=0, sticky="w")

        self.img_out_var = tk.StringVar()
        ttk.Entry(bottom, textvariable=self.img_out_var).grid(row=0, column=1, sticky="ew", padx=8)
        ttk.Button(bottom, text="Escolher...", command=self._pick_img_output).grid(row=0, column=2)

        ttk.Button(parent, text="Converter", command=self._do_images_convert).grid(row=3, column=0, pady=(8, 12))
        self._on_one_pdf_toggle()

    def _add_image_path(self, img_path: Path):
        if img_path.suffix.lower() not in IMG_EXTS or not img_path.exists() or not img_path.is_file():
            return
        rp = self._normalize_path(img_path)
        if rp in self._img_seen:
            return
        self._img_seen.add(rp)
        self.img_list.append(img_path)
        self.img_listbox.insert(tk.END, img_path.name)
        if not self.img_out_var.get().strip():
            self.img_out_var.set(str(img_path.with_name(f"{img_path.stem}_IMAGENS.pdf")) if self.one_pdf_var.get() else str(img_path.parent))

    def _on_drop_images(self, event):
        paths = _parse_dnd_files(event.data)
        added = 0
        for raw in paths:
            p = Path(raw)
            if p.suffix.lower() in IMG_EXTS:
                before = len(self.img_list)
                self._add_image_path(p)
                if len(self.img_list) > before:
                    added += 1
        if added == 0:
            messagebox.showwarning("Nenhuma imagem válida", "Solte arquivos de imagem (png, jpg, tiff, bmp, webp, gif).")
        return "break"

    def _add_images(self):
        paths = filedialog.askopenfilenames(title="Selecione imagens", filetypes=[("Imagens", "*.png *.jpg *.jpeg *.tif *.tiff *.bmp *.webp *.gif"), ("Todos os arquivos", "*.*")])
        for p in paths:
            self._add_image_path(Path(p))

    def _remove_selected_image(self):
        sel = self.img_listbox.curselection()
        if not sel:
            return
        idx = sel[0]
        removed = self.img_list[idx]
        self.img_listbox.delete(idx)
        del self.img_list[idx]
        self._img_seen.discard(self._normalize_path(removed))

    def _clear_images(self):
        self.img_listbox.delete(0, tk.END)
        self.img_list.clear()
        self._img_seen.clear()

    def _img_move_up(self):
        sel = self.img_listbox.curselection()
        if not sel:
            return
        i = sel[0]
        if i <= 0:
            return
        self.img_list[i - 1], self.img_list[i] = self.img_list[i], self.img_list[i - 1]
        txt = self.img_listbox.get(i)
        self.img_listbox.delete(i)
        self.img_listbox.insert(i - 1, txt)
        self.img_listbox.selection_set(i - 1)

    def _img_move_down(self):
        sel = self.img_listbox.curselection()
        if not sel:
            return
        i = sel[0]
        if i >= len(self.img_list) - 1:
            return
        self.img_list[i + 1], self.img_list[i] = self.img_list[i], self.img_list[i + 1]
        txt = self.img_listbox.get(i)
        self.img_listbox.delete(i)
        self.img_listbox.insert(i + 1, txt)
        self.img_listbox.selection_set(i + 1)

    def _on_one_pdf_toggle(self):
        if self.one_pdf_var.get():
            self.img_out_label.configure(text="Salvar PDF final em:")
            if self.img_list and (not self.img_out_var.get().strip() or Path(self.img_out_var.get()).is_dir()):
                self.img_out_var.set(str(self.img_list[0].with_name(f"{self.img_list[0].stem}_IMAGENS.pdf")))
        else:
            self.img_out_label.configure(text="Salvar PDFs na pasta:")
            if self.img_list and (not self.img_out_var.get().strip() or self.img_out_var.get().strip().lower().endswith(".pdf")):
                self.img_out_var.set(str(self.img_list[0].parent))

    def _pick_img_output(self):
        if self.one_pdf_var.get():
            path = filedialog.asksaveasfilename(title="Salvar PDF como", defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
            if path:
                self.img_out_var.set(path)
        else:
            folder = filedialog.askdirectory(title="Escolha a pasta para salvar os PDFs")
            if folder:
                self.img_out_var.set(folder)

    def _do_images_convert(self):
        try:
            if not self.img_list:
                raise ValueError("Adicione imagens na lista (arraste e solte ou use o botão).")
            if not self.img_out_var.get().strip():
                raise ValueError("Escolha o destino de saída.")
            fit_a4 = bool(self.fit_a4_var.get())
            if self.one_pdf_var.get():
                out_pdf = self._unique_outfile(_safe_pdf_path(self.img_out_var.get()))
                convert_images_to_one_pdf(self.img_list, out_pdf, fit_a4=fit_a4)
                messagebox.showinfo("Concluído", f"PDF único gerado com sucesso:\n{out_pdf}")
            else:
                out_dir = Path(self.img_out_var.get().strip().strip('"'))
                out_dir.mkdir(parents=True, exist_ok=True)
                ok = 0
                for img in self.img_list:
                    out_pdf = self._unique_outfile(out_dir / f"{img.stem}.pdf")
                    convert_image_to_pdf_single(img, out_pdf, fit_a4=fit_a4)
                    ok += 1
                messagebox.showinfo("Concluído", f"PDFs gerados com sucesso: {ok}\nPasta:\n{out_dir}")
        except Exception as e:
            messagebox.showerror("Erro", str(e))

    # ===================== Aba 2: Juntar PDFs =====================

    def _build_tab_merge(self, parent: ttk.Frame):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        top = ttk.Frame(parent)
        top.grid(row=0, column=0, sticky="ew", padx=8, pady=(12, 6))
        top.columnconfigure(0, weight=1)
        ttk.Button(top, text="Adicionar PDFs...", command=self._add_pdfs).grid(row=0, column=0, sticky="w")
        ttk.Button(top, text="Remover selecionado", command=self._remove_selected).grid(row=0, column=1, padx=8)
        ttk.Button(top, text="Limpar lista", command=self._clear_list).grid(row=0, column=2)
        mid = ttk.LabelFrame(parent, text="Arraste e solte PDFs aqui (a ordem da lista será a ordem final)")
        mid.grid(row=1, column=0, sticky="nsew", padx=8, pady=6)
        mid.columnconfigure(0, weight=1)
        mid.rowconfigure(0, weight=1)
        self.pdf_list: list[Path] = []
        self.listbox = tk.Listbox(mid, selectmode=tk.SINGLE)
        self.listbox.grid(row=0, column=0, sticky="nsew", padx=(8, 0), pady=8)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self.listbox.yview)
        sb.grid(row=0, column=1, sticky="ns", pady=8)
        self.listbox.configure(yscrollcommand=sb.set)
        buttons = ttk.Frame(mid)
        buttons.grid(row=0, column=2, sticky="ns", padx=10, pady=8)
        ttk.Button(buttons, text="↑ Subir", command=self._move_up).grid(row=0, column=0, pady=(0, 6), sticky="ew")
        ttk.Button(buttons, text="↓ Descer", command=self._move_down).grid(row=1, column=0, pady=(0, 6), sticky="ew")
        for w in (mid, self.listbox):
            try:
                w.drop_target_register(DND_FILES)
                w.dnd_bind("<<Drop>>", self._on_drop_pdfs)
            except Exception:
                pass
        bottom = ttk.Frame(parent)
        bottom.grid(row=2, column=0, sticky="ew", padx=8, pady=(10, 6))
        bottom.columnconfigure(1, weight=1)
        ttk.Label(bottom, text="Salvar PDF final em:").grid(row=0, column=0, sticky="w")
        self.merge_out_var = tk.StringVar()
        ttk.Entry(bottom, textvariable=self.merge_out_var).grid(row=0, column=1, sticky="ew", padx=8)
        ttk.Button(bottom, text="Escolher...", command=self._pick_merge_out).grid(row=0, column=2)
        ttk.Button(parent, text="Juntar PDFs", command=self._do_merge).grid(row=3, column=0, pady=(8, 12))

    def _add_pdf_path(self, pdf_path: Path):
        rp = self._normalize_path(pdf_path)
        if rp in self._merge_seen or not pdf_path.exists() or not pdf_path.is_file():
            return
        self._merge_seen.add(rp)
        self.pdf_list.append(pdf_path)
        self.listbox.insert(tk.END, pdf_path.name)
        if not self.merge_out_var.get().strip():
            self.merge_out_var.set(str(pdf_path.with_name(f"{pdf_path.stem}_JUNTADO.pdf")))

    def _on_drop_pdfs(self, event):
        paths = _parse_dnd_files(event.data)
        for p in paths:
            if p.lower().endswith(".pdf"):
                self._add_pdf_path(Path(p))
        return "break"

    def _add_pdfs(self):
        paths = filedialog.askopenfilenames(title="Selecione PDFs para juntar (ordem importa)", filetypes=[("PDF", "*.pdf")])
        for p in paths:
            self._add_pdf_path(Path(p))

    def _remove_selected(self):
        sel = self.listbox.curselection()
        if not sel:
            return
        idx = sel[0]
        removed = self.pdf_list[idx]
        self.listbox.delete(idx)
        del self.pdf_list[idx]
        self._merge_seen.discard(self._normalize_path(removed))

    def _clear_list(self):
        self.listbox.delete(0, tk.END)
        self.pdf_list.clear()
        self._merge_seen.clear()

    def _move_up(self):
        sel = self.listbox.curselection()
        if not sel:
            return
        i = sel[0]
        if i <= 0:
            return
        self.pdf_list[i - 1], self.pdf_list[i] = self.pdf_list[i], self.pdf_list[i - 1]
        txt = self.listbox.get(i)
        self.listbox.delete(i)
        self.listbox.insert(i - 1, txt)
        self.listbox.selection_set(i - 1)

    def _move_down(self):
        sel = self.listbox.curselection()
        if not sel:
            return
        i = sel[0]
        if i >= len(self.pdf_list) - 1:
            return
        self.pdf_list[i + 1], self.pdf_list[i] = self.pdf_list[i], self.pdf_list[i + 1]
        txt = self.listbox.get(i)
        self.listbox.delete(i)
        self.listbox.insert(i + 1, txt)
        self.listbox.selection_set(i + 1)

    def _pick_merge_out(self):
        path = filedialog.asksaveasfilename(title="Salvar PDF final como", defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if path:
            self.merge_out_var.set(path)

    def _ask_password(self, filename: str) -> str | None:
        return simpledialog.askstring("PDF protegido", f"Informe a senha do PDF:\n{filename}", show="*")

    def _do_merge(self):
        try:
            if not self.merge_out_var.get().strip():
                raise ValueError("Escolha o caminho de saída do PDF final.")
            out_path = self._unique_outfile(_safe_pdf_path(self.merge_out_var.get()))
            merge_pdfs(self.pdf_list, out_path, password_cb=self._ask_password)
            messagebox.showinfo("Concluído", f"PDF juntado com sucesso:\n{out_path}")
        except Exception as e:
            messagebox.showerror("Erro", str(e))

    # ===================== Aba 3: Word/Excel -> PDF =====================

    def _build_tab_office(self, parent: ttk.Frame):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1)

        top = ttk.Frame(parent)
        top.grid(row=0, column=0, sticky="ew", padx=8, pady=(12, 6))
        top.columnconfigure(0, weight=1)
        ttk.Button(top, text="Adicionar Word/Excel...", command=self._add_office_files).grid(row=0, column=0, sticky="w")
        ttk.Button(top, text="Remover selecionado", command=self._remove_selected_office).grid(row=0, column=1, padx=8)
        ttk.Button(top, text="Limpar lista", command=self._clear_office_files).grid(row=0, column=2)

        # Status rápido, visível logo no topo da aba.
        self.office_quick_status_var = tk.StringVar(value="Office: verificação ainda não realizada.")
        ttk.Label(top, textvariable=self.office_quick_status_var).grid(row=1, column=0, columnspan=3, sticky="w", pady=(6, 0))

        mode_box = ttk.LabelFrame(parent, text="Modo de conversão")
        mode_box.grid(row=1, column=0, sticky="ew", padx=8, pady=(0, 6))
        mode_box.columnconfigure(3, weight=1)

        self.office_mode_var = tk.StringVar(value="auto")
        ttk.Radiobutton(mode_box, text="Automático", variable=self.office_mode_var, value="auto").grid(row=0, column=0, sticky="w", padx=8, pady=4)
        ttk.Radiobutton(mode_box, text="Office fiel", variable=self.office_mode_var, value="office").grid(row=0, column=1, sticky="w", padx=8, pady=4)
        ttk.Radiobutton(mode_box, text="Simples sem Office", variable=self.office_mode_var, value="simple").grid(row=0, column=2, sticky="w", padx=8, pady=4)
        ttk.Button(mode_box, text="Verificar Office", command=self._refresh_office_status).grid(row=0, column=4, sticky="e", padx=8, pady=4)

        self.office_detect_var = tk.StringVar(value="Verificação do Office ainda não realizada.")
        ttk.Label(mode_box, textvariable=self.office_detect_var).grid(row=1, column=0, columnspan=5, sticky="w", padx=8, pady=(0, 4))

        self.office_mode_info_var = tk.StringVar(
            value="Automático: usa Word/Excel quando encontrados; caso contrário tenta o modo simples. O modo simples pode perder formatação."
        )
        ttk.Label(mode_box, textvariable=self.office_mode_info_var, wraplength=880).grid(row=2, column=0, columnspan=5, sticky="w", padx=8, pady=(0, 6))

        mid = ttk.LabelFrame(parent, text="Arraste e solte arquivos Word/Excel aqui")
        mid.grid(row=2, column=0, sticky="nsew", padx=8, pady=6)
        mid.columnconfigure(0, weight=1)
        mid.rowconfigure(0, weight=1)

        self.office_list: list[Path] = []
        self.office_listbox = tk.Listbox(mid, selectmode=tk.SINGLE)
        self.office_listbox.grid(row=0, column=0, sticky="nsew", padx=(8, 0), pady=8)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self.office_listbox.yview)
        sb.grid(row=0, column=1, sticky="ns", pady=8)
        self.office_listbox.configure(yscrollcommand=sb.set)
        side = ttk.Frame(mid)
        side.grid(row=0, column=2, sticky="ns", padx=10, pady=8)
        ttk.Button(side, text="↑ Subir", command=self._office_move_up).grid(row=0, column=0, pady=(0, 6), sticky="ew")
        ttk.Button(side, text="↓ Descer", command=self._office_move_down).grid(row=1, column=0, pady=(0, 6), sticky="ew")
        for w in (mid, self.office_listbox):
            try:
                w.drop_target_register(DND_FILES)
                w.dnd_bind("<<Drop>>", self._on_drop_office_files)
            except Exception:
                pass

        bottom = ttk.Frame(parent)
        bottom.grid(row=3, column=0, sticky="ew", padx=8, pady=(10, 6))
        bottom.columnconfigure(1, weight=1)
        ttk.Label(bottom, text="Salvar PDFs na pasta:").grid(row=0, column=0, sticky="w")
        self.office_out_var = tk.StringVar()
        ttk.Entry(bottom, textvariable=self.office_out_var).grid(row=0, column=1, sticky="ew", padx=8)
        ttk.Button(bottom, text="Escolher...", command=self._pick_office_out_dir).grid(row=0, column=2)

        self.office_status_var = tk.StringVar(value="Pronto.")
        ttk.Label(parent, textvariable=self.office_status_var).grid(row=4, column=0, sticky="w", padx=10, pady=(4, 0))
        ttk.Button(parent, text="Converter Word/Excel para PDF", command=self._do_office_convert).grid(row=5, column=0, pady=(8, 12))

        # Faz uma primeira verificação sem bloquear a criação da interface por muito tempo.
        self.after(300, self._refresh_office_status)

    def _set_office_detect_text(self, text: str):
        if hasattr(self, "office_detect_var"):
            self.office_detect_var.set(text)
        if hasattr(self, "office_quick_status_var"):
            self.office_quick_status_var.set(text)

    def _refresh_office_status(self):
        try:
            self._set_office_detect_text("Verificando Microsoft Word/Excel...")
            self.update_idletasks()
            self._office_status = detect_ms_office()

            if not self._office_status.get("platform_ok"):
                self._set_office_detect_text("Office fiel: indisponível. Motivo: automação do Office exige Windows.")
                return
            if not self._office_status.get("pywin32_ok"):
                self._set_office_detect_text("Office fiel: indisponível. Motivo: pywin32 não carregou.")
                return

            word_txt = "Word encontrado" if self._office_status.get("word_ok") else "Word não encontrado"
            excel_txt = "Excel encontrado" if self._office_status.get("excel_ok") else "Excel não encontrado"
            self._set_office_detect_text(f"Office fiel: {word_txt}; {excel_txt}. Modo simples continua disponível para DOCX/DOCM/XLSX/XLSM/CSV.")
        except Exception as e:
            self._set_office_detect_text(f"Falha ao verificar Office: {e}")

    def _add_office_path(self, file_path: Path):
        if file_path.suffix.lower() not in OFFICE_EXTS or not file_path.exists() or not file_path.is_file():
            return
        rp = self._normalize_path(file_path)
        if rp in self._office_seen:
            return
        self._office_seen.add(rp)
        self.office_list.append(file_path)
        self.office_listbox.insert(tk.END, file_path.name)
        if not self.office_out_var.get().strip():
            self.office_out_var.set(str(file_path.parent))

    def _on_drop_office_files(self, event):
        paths = _parse_dnd_files(event.data)
        added = 0
        for raw in paths:
            p = Path(raw)
            if p.suffix.lower() in OFFICE_EXTS:
                before = len(self.office_list)
                self._add_office_path(p)
                if len(self.office_list) > before:
                    added += 1
        if added == 0:
            messagebox.showwarning("Nenhum arquivo válido", "Solte arquivos Word/Excel: doc, docx, docm, rtf, xls, xlsx, xlsm, xlsb ou csv.")
        return "break"

    def _add_office_files(self):
        paths = filedialog.askopenfilenames(
            title="Selecione arquivos Word/Excel",
            filetypes=[
                ("Word/Excel", "*.doc *.docx *.docm *.rtf *.xls *.xlsx *.xlsm *.xlsb *.csv"),
                ("Word", "*.doc *.docx *.docm *.rtf"),
                ("Excel", "*.xls *.xlsx *.xlsm *.xlsb *.csv"),
                ("Todos os arquivos", "*.*"),
            ],
        )
        for p in paths:
            self._add_office_path(Path(p))

    def _remove_selected_office(self):
        sel = self.office_listbox.curselection()
        if not sel:
            return
        idx = sel[0]
        removed = self.office_list[idx]
        self.office_listbox.delete(idx)
        del self.office_list[idx]
        self._office_seen.discard(self._normalize_path(removed))

    def _clear_office_files(self):
        self.office_listbox.delete(0, tk.END)
        self.office_list.clear()
        self._office_seen.clear()

    def _office_move_up(self):
        sel = self.office_listbox.curselection()
        if not sel:
            return
        i = sel[0]
        if i <= 0:
            return
        self.office_list[i - 1], self.office_list[i] = self.office_list[i], self.office_list[i - 1]
        txt = self.office_listbox.get(i)
        self.office_listbox.delete(i)
        self.office_listbox.insert(i - 1, txt)
        self.office_listbox.selection_set(i - 1)

    def _office_move_down(self):
        sel = self.office_listbox.curselection()
        if not sel:
            return
        i = sel[0]
        if i >= len(self.office_list) - 1:
            return
        self.office_list[i + 1], self.office_list[i] = self.office_list[i], self.office_list[i + 1]
        txt = self.office_listbox.get(i)
        self.office_listbox.delete(i)
        self.office_listbox.insert(i + 1, txt)
        self.office_listbox.selection_set(i + 1)

    def _pick_office_out_dir(self):
        folder = filedialog.askdirectory(title="Escolha a pasta para salvar os PDFs")
        if folder:
            self.office_out_var.set(folder)

    def _set_office_status(self, text: str):
        self.office_status_var.set(text)
        self.update_idletasks()

    def _do_office_convert(self):
        try:
            if not self.office_list:
                raise ValueError("Adicione arquivos Word/Excel na lista.")
            if not self.office_out_var.get().strip():
                raise ValueError("Escolha a pasta de saída.")

            mode = self.office_mode_var.get()
            out_dir = Path(self.office_out_var.get().strip().strip('"'))

            # Atualiza status antes da conversão automática/Office.
            if mode in {"auto", "office"}:
                self._refresh_office_status()

            if mode == "office":
                missing = []
                if any(p.suffix.lower() in WORD_EXTS for p in self.office_list) and not self._office_status.get("word_ok"):
                    missing.append("Microsoft Word")
                if any(p.suffix.lower() in EXCEL_EXTS for p in self.office_list) and not self._office_status.get("excel_ok"):
                    missing.append("Microsoft Excel")
                if missing:
                    raise RuntimeError(
                        "Modo Office fiel selecionado, mas não encontrei: " + ", ".join(missing) + ".\n"
                        "Use o modo 'Simples sem Office' ou instale o Office correspondente."
                    )

            converted, errors = convert_office_files_to_pdf(
                self.office_list,
                out_dir,
                unique_cb=self._unique_outfile,
                status_cb=self._set_office_status,
                mode=mode,
                office_status=self._office_status,
            )

            if errors:
                msg = f"Convertidos com sucesso: {len(converted)}\nErros: {len(errors)}\n\n" + "\n".join(errors[:8])
                if len(errors) > 8:
                    msg += f"\n... e mais {len(errors) - 8} erro(s)."
                messagebox.showwarning("Concluído com erros", msg)
            else:
                messagebox.showinfo("Concluído", f"Arquivos convertidos com sucesso: {len(converted)}\nPasta:\n{out_dir}")

        except Exception as e:
            self._set_office_status("Erro na conversão.")
            messagebox.showerror("Erro", str(e))


if __name__ == "__main__":
    App().mainloop()
